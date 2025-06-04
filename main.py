from fastapi import FastAPI, HTTPException, BackgroundTasks, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional
import uvicorn
from datetime import datetime
import json
import uuid
import os

from models import (
    Product, ChatRequest, ChatResponse, ProductQuery, 
    ProductRecommendation, ChatMessage, FileUpload, FileUploadResponse
)
from chatbot_service import ChatbotService
from vector_store import VectorStore
from config import settings

# Initialize FastAPI app
app = FastAPI(
    title="Customer Service Chatbot API",
    description="A backend API for customer service chatbot with product recommendations",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize services
chatbot_service = ChatbotService()
vector_store = VectorStore()

# File upload tracking
uploaded_files: List[FileUpload] = []
UPLOAD_DIR = "uploads"

# Create upload directory if it doesn't exist
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.on_event("startup")
async def startup_event():
    """Initialize the application on startup"""
    print("🚀 Starting Customer Service Chatbot API...")
    print(f"OpenAI Model: {settings.OPENAI_MODEL}")
    print(f"Embedding Model: {settings.EMBEDDING_MODEL}")
    print(f"Chroma DB Path: {settings.CHROMA_PERSIST_DIRECTORY}")
    
    # Auto-populate sample data if database is empty
    try:
        existing_products = vector_store.get_all_products()
        if len(existing_products) == 0:
            print("📚 Database is empty. Loading sample data...")
            await populate_sample_data()
            print("✅ Sample data loaded successfully!")
        else:
            print(f"📊 Found {len(existing_products)} existing products in database")
    except Exception as e:
        print(f"⚠️ Error checking/loading sample data: {e}")

async def populate_sample_data():
    """Populate the database with sample products"""
    from sample_data import SAMPLE_PRODUCTS
    
    success_count = 0
    for product in SAMPLE_PRODUCTS:
        try:
            if vector_store.add_product(product):
                success_count += 1
        except Exception as e:
            print(f"❌ Failed to add {product.name}: {e}")
    
    print(f"📦 Successfully loaded {success_count}/{len(SAMPLE_PRODUCTS)} sample products")

@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "message": "Customer Service Chatbot API",
        "status": "running",
        "timestamp": datetime.now().isoformat()
    }

# Chat Endpoints
@app.post("/chat", response_model=ChatResponse)
async def chat_with_bot(request: ChatRequest):
    """Chat with the customer service chatbot"""
    try:
        response = await chatbot_service.chat(request)
        return response
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing chat: {str(e)}")

@app.get("/chat/history/{conversation_id}")
async def get_chat_history(conversation_id: str):
    """Get conversation history"""
    try:
        history = chatbot_service.get_conversation_history(conversation_id)
        return {"conversation_id": conversation_id, "messages": history}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving history: {str(e)}")

@app.delete("/chat/{conversation_id}")
async def clear_conversation(conversation_id: str):
    """Clear a conversation"""
    success = chatbot_service.clear_conversation(conversation_id)
    if success:
        return {"message": "Conversation cleared successfully"}
    else:
        raise HTTPException(status_code=404, detail="Conversation not found")

# Product Management Endpoints
@app.post("/products")
async def add_product(product: Product, background_tasks: BackgroundTasks):
    """Add a new product to the catalog"""
    try:
        # Add product to vector store in background
        success = vector_store.add_product(product)
        if success:
            return {"message": "Product added successfully", "product_id": product.id}
        else:
            raise HTTPException(status_code=500, detail="Failed to add product to vector store")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error adding product: {str(e)}")

@app.get("/products", response_model=List[Product])
async def get_all_products():
    """Get all products from the catalog"""
    try:
        products = vector_store.get_all_products()
        return products
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving products: {str(e)}")

@app.put("/products/{product_id}")
async def update_product(product_id: str, product: Product, background_tasks: BackgroundTasks):
    """Update an existing product"""
    try:
        if product.id != product_id:
            raise HTTPException(status_code=400, detail="Product ID mismatch")
        
        # Update product in vector store
        success = vector_store.update_product(product)
        if success:
            return {"message": "Product updated successfully"}
        else:
            raise HTTPException(status_code=500, detail="Failed to update product")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating product: {str(e)}")

@app.delete("/products/{product_id}")
async def delete_product(product_id: str):
    """Delete a product from the catalog"""
    try:
        success = vector_store.delete_product(product_id)
        if success:
            return {"message": "Product deleted successfully"}
        else:
            raise HTTPException(status_code=404, detail="Product not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting product: {str(e)}")

# Product Search and Recommendation Endpoints
@app.post("/products/search")
async def search_products(query: ProductQuery):
    """Search for products based on a query"""
    try:
        if query.category:
            results = vector_store.search_by_category(query.category, query.limit)
        else:
            results = vector_store.search_products(query.query, query.limit)
        
        # Filter by price range if specified
        if query.price_range:
            min_price, max_price = query.price_range
            results = [
                r for r in results 
                if min_price <= r["product"].price <= max_price
            ]
        
        return {
            "query": query.query,
            "results": [
                {
                    "product": result["product"],
                    "similarity_score": result["similarity_score"]
                }
                for result in results
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error searching products: {str(e)}")

@app.post("/recommendations", response_model=ProductRecommendation)
async def get_recommendations(query: str, limit: int = 5):
    """Get product recommendations with reasoning"""
    try:
        recommendations = await chatbot_service.get_product_recommendations(query, limit)
        return recommendations
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting recommendations: {str(e)}")

@app.get("/categories")
async def get_categories():
    """Get all available product categories"""
    try:
        products = vector_store.get_all_products()
        categories = list(set(product.category for product in products))
        return {"categories": sorted(categories)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving categories: {str(e)}")

# Utility Endpoints
@app.post("/products/bulk")
async def add_products_bulk(products: List[Product]):
    """Add multiple products in bulk"""
    try:
        success_count = 0
        for product in products:
            if vector_store.add_product(product):
                success_count += 1
        
        return {
            "message": f"Successfully added {success_count}/{len(products)} products",
            "success_count": success_count,
            "total_count": len(products)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error adding products in bulk: {str(e)}")

@app.get("/stats")
async def get_stats():
    """Get API statistics"""
    try:
        all_products = vector_store.get_all_products()
        categories = list(set(product.category for product in all_products))
        available_products = [p for p in all_products if p.availability]
        
        return {
            "total_products": len(all_products),
            "available_products": len(available_products),
            "categories": len(categories),
            "conversations_active": chatbot_service.get_active_conversations_count(),
            "category_list": sorted(categories)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving stats: {str(e)}")

# File Upload Endpoints
@app.post("/upload/products", response_model=FileUploadResponse)
async def upload_product_file(file: UploadFile = File(...)):
    """Upload a JSON, CSV, PDF, TXT, or other text-based file containing product data"""
    try:
        # Validate file type - support multiple formats
        supported_extensions = ('.json', '.csv', '.pdf', '.txt', '.md', '.docx', '.xml')
        if not file.filename.lower().endswith(supported_extensions):
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Supported formats: {', '.join(supported_extensions)}"
            )
        
        # Save uploaded file
        upload_id = str(uuid.uuid4())
        file_path = os.path.join(UPLOAD_DIR, f"{upload_id}_{file.filename}")
        
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        # Parse and add products based on file type
        products_added = 0
        file_extension = file.filename.lower().split('.')[-1]
        
        if file_extension == 'json':
            products_added = await _process_json_file(file_path)
        elif file_extension == 'csv':
            products_added = await _process_csv_file(file_path)
        elif file_extension == 'pdf':
            products_added = await _process_pdf_file(file_path)
        elif file_extension in ['txt', 'md']:
            products_added = await _process_text_file(file_path)
        elif file_extension == 'docx':
            products_added = await _process_docx_file(file_path)
        elif file_extension == 'xml':
            products_added = await _process_xml_file(file_path)
        
        # Track uploaded file
        file_upload = FileUpload(
            filename=file.filename,
            upload_time=datetime.now(),
            file_size=len(content),
            products_added=products_added,
            status="success" if products_added > 0 else "failed"
        )
        uploaded_files.append(file_upload)
        
        return FileUploadResponse(
            message=f"Successfully uploaded {file.filename} and added {products_added} products",
            filename=file.filename,
            products_added=products_added,
            upload_id=upload_id
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading file: {str(e)}")

@app.get("/uploads", response_model=List[FileUpload])
async def get_uploaded_files():
    """Get list of all uploaded files"""
    return uploaded_files

@app.get("/uploads/{filename}")
async def get_upload_details(filename: str):
    """Get details of a specific uploaded file"""
    for upload in uploaded_files:
        if upload.filename == filename:
            return upload
    raise HTTPException(status_code=404, detail="File not found")

async def _process_json_file(file_path: str) -> int:
    """Process a JSON file containing product data"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        products_added = 0
        
        # Handle both single product and array of products
        if isinstance(data, list):
            for product_data in data:
                try:
                    product = Product(**product_data)
                    if vector_store.add_product(product):
                        products_added += 1
                except Exception as e:
                    print(f"Error adding product: {e}")
        elif isinstance(data, dict):
            try:
                product = Product(**data)
                if vector_store.add_product(product):
                    products_added += 1
            except Exception as e:
                print(f"Error adding product: {e}")
        
        return products_added
        
    except Exception as e:
        print(f"Error processing JSON file: {e}")
        return 0

async def _process_csv_file(file_path: str) -> int:
    """Process a CSV file containing product data"""
    try:
        import pandas as pd
        
        df = pd.read_csv(file_path)
        products_added = 0
        
        for _, row in df.iterrows():
            try:
                # Parse features and specifications from string format
                features = row.get('features', '').split(',') if row.get('features') else []
                features = [f.strip() for f in features if f.strip()]
                
                specs = {}
                if 'specifications' in row and pd.notna(row['specifications']):
                    # Assume specifications are in JSON format in the CSV
                    specs = json.loads(row['specifications'])
                
                product = Product(
                    id=row['id'],
                    name=row['name'],
                    description=row['description'],
                    category=row['category'],
                    price=float(row['price']),
                    features=features,
                    specifications=specs,
                    availability=bool(row.get('availability', True))
                )
                
                if vector_store.add_product(product):
                    products_added += 1
                    
            except Exception as e:
                print(f"Error adding product from CSV row: {e}")
        
        return products_added
        
    except Exception as e:
        print(f"Error processing CSV file: {e}")
        return 0

async def _process_pdf_file(file_path: str) -> int:
    """Process a PDF file containing product data"""
    try:
        import PyPDF2
        import pdfplumber
        import re
        
        # Try pdfplumber first (better text extraction)
        text_content = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
        except Exception:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
        
        if not text_content.strip():
            print("No text content extracted from PDF")
            return 0
        
        # Parse text content for product information
        return await _parse_text_for_products(text_content)
        
    except Exception as e:
        print(f"Error processing PDF file: {e}")
        return 0

async def _process_text_file(file_path: str) -> int:
    """Process a plain text or markdown file containing product data"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        return await _parse_text_for_products(text_content)
        
    except Exception as e:
        print(f"Error processing text file: {e}")
        return 0

async def _process_docx_file(file_path: str) -> int:
    """Process a Word document containing product data"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_content = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + "\t"
                text_content += "\n"
        
        return await _parse_text_for_products(text_content)
        
    except Exception as e:
        print(f"Error processing DOCX file: {e}")
        return 0

async def _process_xml_file(file_path: str) -> int:
    """Process an XML file containing product data"""
    try:
        import xml.etree.ElementTree as ET
        
        tree = ET.parse(file_path)
        root = tree.getroot()
        products_added = 0
        
        # Try to parse as structured XML first
        products = root.findall('.//product') or root.findall('.//item') or [root]
        
        for product_elem in products:
            try:
                # Extract basic product info
                product_data = {}
                
                # Common field mappings
                field_mappings = {
                    'id': ['id', 'product_id', 'sku'],
                    'name': ['name', 'title', 'product_name'],
                    'description': ['description', 'desc', 'summary'],
                    'category': ['category', 'type', 'class'],
                    'price': ['price', 'cost', 'amount']
                }
                
                for field, possible_names in field_mappings.items():
                    for name in possible_names:
                        elem = product_elem.find(f'.//{name}') or product_elem.find(name)
                        if elem is not None and elem.text:
                            if field == 'price':
                                product_data[field] = float(elem.text.replace('$', '').replace(',', ''))
                            else:
                                product_data[field] = elem.text.strip()
                            break
                
                # Extract features
                features = []
                features_elem = product_elem.find('.//features') or product_elem.find('features')
                if features_elem is not None:
                    for feature in features_elem.findall('.//feature') or features_elem.findall('item'):
                        if feature.text:
                            features.append(feature.text.strip())
                
                # Extract specifications
                specs = {}
                specs_elem = product_elem.find('.//specifications') or product_elem.find('specs')
                if specs_elem is not None:
                    for spec in specs_elem:
                        if spec.text:
                            specs[spec.tag] = spec.text.strip()
                
                # Create product if we have minimum required fields
                if all(field in product_data for field in ['id', 'name', 'description', 'category', 'price']):
                    product = Product(
                        id=product_data['id'],
                        name=product_data['name'],
                        description=product_data['description'],
                        category=product_data['category'],
                        price=product_data['price'],
                        features=features or [],
                        specifications=specs or {},
                        availability=True
                    )
                    
                    if vector_store.add_product(product):
                        products_added += 1
                        
            except Exception as e:
                print(f"Error processing XML product: {e}")
        
        # If structured parsing failed, try text parsing
        if products_added == 0:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            return await _parse_text_for_products(text_content)
        
        return products_added
        
    except Exception as e:
        print(f"Error processing XML file: {e}")
        return 0

async def _parse_text_for_products(text_content: str) -> int:
    """Parse unstructured text content for product information using AI"""
    try:
        # Use OpenAI to extract product information from text
        from openai import OpenAI
        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        
        prompt = f"""
        Extract product information from the following text and format it as a JSON array of products.
        Each product should have these fields:
        - id: a unique identifier (generate if not present)
        - name: product name
        - description: product description
        - category: product category (e.g., "Laptop", "Smartphone", "Monitor", "Aksesoris")
        - price: price as a number (convert currencies to IDR if needed, assume USD if currency not specified)
        - features: array of key features
        - specifications: object with technical specifications
        - availability: boolean (assume true if not specified)

        Text content:
        {text_content[:4000]}  # Limit to avoid token limits

        Return ONLY the JSON array, no other text:
        """
        
        response = client.chat.completions.create(
            model=settings.OPENAI_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000,
            temperature=0.1
        )
        
        # Parse the AI response as JSON
        ai_response = response.choices[0].message.content.strip()
        
        # Clean up the response (remove markdown formatting if present)
        if ai_response.startswith('```json'):
            ai_response = ai_response[7:]
        if ai_response.endswith('```'):
            ai_response = ai_response[:-3]
        
        products_data = json.loads(ai_response)
        
        # Add products to the database
        products_added = 0
        if isinstance(products_data, list):
            for product_data in products_data:
                try:
                    # Ensure we have an ID
                    if 'id' not in product_data or not product_data['id']:
                        product_data['id'] = f"extracted-{uuid.uuid4().hex[:8]}"
                    
                    product = Product(**product_data)
                    if vector_store.add_product(product):
                        products_added += 1
                except Exception as e:
                    print(f"Error creating product from AI extraction: {e}")
        
        return products_added
        
    except Exception as e:
        print(f"Error parsing text for products: {e}")
        return 0

if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    ) 